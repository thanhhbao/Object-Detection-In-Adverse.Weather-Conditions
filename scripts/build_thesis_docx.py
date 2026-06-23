#!/usr/bin/env python3
"""Sinh khóa luận tốt nghiệp (.docx) đúng quy định trình bày của ĐH Nguyễn Tất Thành.

Chạy:  python scripts/build_thesis_docx.py
Kết quả: docs/thesis/Khoa_luan_tot_nghiep.docx

Script áp sẵn: lề, font Times New Roman 13, giãn dòng 1.5, tiêu đề chương 16 /
mục 14, đánh số trang La Mã (phần đầu) -> Ả Rập (từ Chương 1), mục lục tự động,
trang bìa + bìa trong. Nội dung học thuật điền trực tiếp trong CONTENT bên dưới.

Thông tin cá nhân sửa trong INFO (hoặc sửa thẳng trong Word sau khi mở).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "thesis" / "Khoa_luan_tot_nghiep.docx"

INFO = {
    "title": "THIẾT KẾ VÀ PHÁT TRIỂN MÔ HÌNH HỌC SÂU PHÁT HIỆN NGƯỜI VÀ "
    "PHƯƠNG TIỆN TRONG ĐIỀU KIỆN THỜI TIẾT BẤT LỢI PHỤC VỤ GIÁM SÁT "
    "GIAO THÔNG THÔNG MINH",
    "gvhd": "(Họ tên giảng viên hướng dẫn)",
    "sv": "(Họ tên sinh viên)",
    "mssv": "(MSSV)",
    "khoa": "(Khóa)",
    "nganh": "Công nghệ thông tin",
    "thoi_gian": "TP. HCM, tháng … năm 2026",
}


# ---------------------------------------------------------------------------
# Tiện ích định dạng (XML thấp cấp cho những thứ python-docx không bọc sẵn)
# ---------------------------------------------------------------------------


def set_cell_or_section_page_numbering(section, fmt: str, start: int | None = None) -> None:
    sectPr = section._sectPr
    for existing in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(existing)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)  # "lowerRoman" hoặc "decimal"
    if start is not None:
        pg.set(qn("w:start"), str(start))
    sectPr.append(pg)


def add_footer_page_number(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instr); run._r.append(end)


def add_toc(document, instruction: str, placeholder: str) -> None:
    para = document.add_paragraph()
    run = para.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = placeholder
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, sep, text, end):
        run._r.append(node)


def configure_styles(document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    heading_specs = {"Heading 1": 16, "Heading 2": 14, "Heading 3": 13}
    for name, size in heading_specs.items():
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def configure_margins(section) -> None:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)


# ---------------------------------------------------------------------------
# Khối nội dung tiện dụng
# ---------------------------------------------------------------------------


def centered(document, text, *, bold=False, size=13, caps=False, space_before=0):
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(space_before)
    run = para.add_run(text.upper() if caps else text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return para


def body(document, text):
    para = document.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Cm(1.0)
    return para


def chapter(document, text):
    document.add_paragraph(text, style="Heading 1").alignment = WD_ALIGN_PARAGRAPH.CENTER


def section_heading(document, text):
    document.add_paragraph(text, style="Heading 2")


def sub_heading(document, text):
    document.add_paragraph(text, style="Heading 3")


def front_title(document, text):
    """Tiêu đề trang đầu (Lời cảm ơn, Mục lục...) — KHÔNG dùng Heading để khỏi
    lọt vào mục lục tự động; bold, in hoa, size 14, canh giữa."""
    centered(document, text, bold=True, size=14, caps=True, space_before=6)


# ---------------------------------------------------------------------------
# Trang bìa và phần đầu
# ---------------------------------------------------------------------------


def build_cover(document) -> None:
    centered(document, "BỘ GIÁO DỤC VÀ ĐÀO TẠO", bold=True, size=14)
    centered(document, "TRƯỜNG ĐẠI HỌC NGUYỄN TẤT THÀNH", bold=True, size=16)
    centered(document, "KHOA CÔNG NGHỆ THÔNG TIN", bold=True, size=16)
    centered(document, "[ CHÈN LOGO TRƯỜNG ]", size=12, space_before=24)
    centered(document, "KHÓA LUẬN TỐT NGHIỆP", bold=True, size=16, space_before=36)
    centered(document, INFO["title"], bold=True, size=20, space_before=24)
    document.add_paragraph().paragraph_format.space_after = Pt(48)
    for label, key in (
        ("Giảng viên hướng dẫn: ", "gvhd"),
        ("Sinh viên thực hiện: ", "sv"),
        ("MSSV: ", "mssv"),
        ("Khóa: ", "khoa"),
        ("Ngành/chuyên ngành: ", "nganh"),
    ):
        para = document.add_paragraph()
        run = para.add_run(label + INFO[key])
        run.bold = True
        run.font.size = Pt(14)
    centered(document, INFO["thoi_gian"], bold=True, size=13, space_before=36)


def build_front_matter(document) -> None:
    document.add_page_break()
    front_title(document, "LỜI CẢM ƠN")
    body(document, "(Sinh viên viết lời cảm ơn tới giảng viên hướng dẫn, "
         "nhà trường, gia đình và bạn bè...)")

    document.add_page_break()
    front_title(document, "LỜI MỞ ĐẦU")
    for paragraph in CONTENT["loi_mo_dau"]:
        body(document, paragraph)

    document.add_page_break()
    front_title(document, "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN")
    body(document, "(Trang dành cho giảng viên hướng dẫn nhận xét.)")

    document.add_page_break()
    front_title(document, "MỤC LỤC")
    add_toc(document, 'TOC \\o "1-3" \\h \\z \\u',
            "Nhấn vào đây, bấm Ctrl+A rồi F9 để cập nhật mục lục.")

    document.add_page_break()
    front_title(document, "DANH MỤC BẢNG")
    add_toc(document, 'TOC \\h \\z \\c "Bảng"', "Cập nhật bằng F9 sau khi chèn bảng có caption.")

    document.add_page_break()
    front_title(document, "DANH MỤC HÌNH")
    add_toc(document, 'TOC \\h \\z \\c "Hình"', "Cập nhật bằng F9 sau khi chèn hình có caption.")

    document.add_page_break()
    front_title(document, "DANH MỤC CHỮ VIẾT TẮT")
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].paragraphs[0].add_run("Chữ viết tắt").bold = True
    table.rows[0].cells[1].paragraphs[0].add_run("Ý nghĩa").bold = True
    for abbr, meaning in CONTENT["abbreviations"]:
        cells = table.add_row().cells
        cells[0].text = abbr
        cells[1].text = meaning


# ---------------------------------------------------------------------------
# Nội dung học thuật (Chương 1 viết đầy đủ; Chương 2–5 dựng khung, điền sau)
# ---------------------------------------------------------------------------

CONTENT = {
    "loi_mo_dau": [
        "Giám sát giao thông thông minh ngày càng đóng vai trò quan trọng trong "
        "quản lý đô thị, an toàn giao thông và điều phối phương tiện. Một thành "
        "phần cốt lõi của các hệ thống này là khả năng phát hiện chính xác người "
        "và phương tiện từ hình ảnh camera. Tuy nhiên, trong điều kiện thời tiết "
        "bất lợi như sương mù, mưa, bão cát hay tuyết, chất lượng hình ảnh suy "
        "giảm mạnh khiến độ chính xác của các mô hình phát hiện giảm sút.",
        "Khóa luận này tập trung thiết kế và phát triển mô hình học sâu phát hiện "
        "người và phương tiện trong điều kiện thời tiết bất lợi, đồng thời so sánh "
        "có hệ thống các hướng tiếp cận hiện đại để tìm ra giải pháp cân bằng giữa "
        "độ chính xác và tốc độ, phù hợp cho ứng dụng giám sát giao thông thực tế.",
    ],
    "abbreviations": [
        ("AP", "Average Precision"),
        ("BDD100K", "Berkeley DeepDrive 100K dataset"),
        ("CBAM", "Convolutional Block Attention Module"),
        ("CNN", "Convolutional Neural Network"),
        ("DAWN", "Detection in Adverse Weather Nature dataset"),
        ("FPN", "Feature Pyramid Network"),
        ("FPS", "Frames Per Second"),
        ("mAP", "mean Average Precision"),
        ("RT-DETR", "Real-Time Detection Transformer"),
        ("YOLO", "You Only Look Once"),
    ],
}


def build_chapter1(document) -> None:
    chapter(document, "CHƯƠNG 1. GIỚI THIỆU CHUNG")

    section_heading(document, "1.1. Lý do chọn đề tài")
    body(document, "Hệ thống giám sát giao thông thông minh dựa trên thị giác máy "
         "tính đang được triển khai rộng rãi nhằm giám sát mật độ, phát hiện vi "
         "phạm và hỗ trợ điều phối giao thông. Phát hiện người và phương tiện là "
         "bài toán nền tảng của các hệ thống này.")
    body(document, "Trong thực tế, camera giám sát phải hoạt động ở mọi điều kiện "
         "thời tiết. Sương mù làm giảm độ tương phản, mưa gây nhiễu và che khuất, "
         "bão cát và tuyết làm biến đổi màu sắc và che lấp đối tượng. Những yếu tố "
         "này khiến các mô hình phát hiện được huấn luyện trên dữ liệu thời tiết "
         "tốt suy giảm đáng kể về độ chính xác. Do đó, nghiên cứu mô hình phát "
         "hiện bền vững với thời tiết bất lợi có ý nghĩa thực tiễn cao.")

    section_heading(document, "1.2. Mục tiêu nghiên cứu")
    body(document, "Khóa luận hướng tới các mục tiêu: (1) xây dựng quy trình huấn "
         "luyện mô hình phát hiện người và phương tiện thích nghi với thời tiết "
         "bất lợi; (2) so sánh có hệ thống ba hướng tiếp cận tiêu biểu là mô hình "
         "một giai đoạn (YOLO), mô hình hai giai đoạn (Faster R-CNN) và mô hình "
         "dựa trên transformer (RT-DETR); (3) khảo sát tác động của cơ chế chú ý "
         "CBAM đối với hiệu năng phát hiện trong điều kiện thời tiết khó.")

    section_heading(document, "1.3. Đối tượng và phạm vi nghiên cứu")
    body(document, "Đối tượng nghiên cứu là sáu lớp đối tượng giao thông: người "
         "(person), xe đạp (bicycle), ô tô (car), xe máy (motorcycle), xe buýt "
         "(bus) và xe tải (truck).")
    body(document, "Phạm vi nghiên cứu sử dụng hai bộ dữ liệu công khai: BDD100K "
         "cho giai đoạn thích nghi miền lái xe và DAWN cho giai đoạn thời tiết "
         "bất lợi. Toàn bộ thực nghiệm được thực hiện trên nền tảng Google Colab "
         "với GPU NVIDIA Tesla T4.")

    section_heading(document, "1.4. Phương pháp nghiên cứu")
    body(document, "Khóa luận áp dụng chiến lược học chuyển giao lũy tiến: khởi "
         "tạo từ trọng số huấn luyện trên COCO, tinh chỉnh trên BDD100K để học "
         "miền lái xe, sau đó tinh chỉnh tiếp trên DAWN để thích nghi thời tiết "
         "bất lợi. Các mô hình được đánh giá bằng các độ đo chuẩn (Precision, "
         "Recall, mAP50, mAP50-95), tốc độ suy luận (FPS) và số lượng tham số, "
         "kèm phân tích theo từng lớp và theo từng điều kiện thời tiết.")

    section_heading(document, "1.5. Bố cục khóa luận")
    body(document, "Khóa luận gồm năm chương. Chương 1 giới thiệu chung. Chương 2 "
         "trình bày cơ sở lý luận và các nghiên cứu liên quan. Chương 3 mô tả mô "
         "hình lý thuyết và giải pháp đề xuất. Chương 4 trình bày mô hình thực "
         "nghiệm, kết quả và phân tích. Chương 5 kết luận và đề xuất hướng phát "
         "triển.")


def build_remaining_chapters(document) -> None:
    for title in (
        "CHƯƠNG 2. CƠ SỞ LÝ LUẬN",
        "CHƯƠNG 3. MÔ HÌNH LÝ THUYẾT",
        "CHƯƠNG 4. MÔ HÌNH THỰC NGHIỆM",
        "CHƯƠNG 5. KẾT LUẬN VÀ KIẾN NGHỊ",
    ):
        document.add_page_break()
        chapter(document, title)
        body(document, "(Nội dung đang được biên soạn.)")


# ---------------------------------------------------------------------------
# Lắp ráp tài liệu
# ---------------------------------------------------------------------------


def main() -> None:
    document = Document()
    configure_styles(document)

    front = document.sections[0]
    configure_margins(front)
    set_cell_or_section_page_numbering(front, "lowerRoman", start=1)
    add_footer_page_number(front)

    build_cover(document)
    build_front_matter(document)

    # Section break: từ Chương 1 đánh số trang Ả Rập bắt đầu từ 1.
    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_margins(body_section)
    set_cell_or_section_page_numbering(body_section, "decimal", start=1)
    add_footer_page_number(body_section)

    build_chapter1(document)
    build_remaining_chapters(document)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
