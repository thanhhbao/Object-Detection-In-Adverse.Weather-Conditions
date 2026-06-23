#!/usr/bin/env python3
"""Dọn file mẫu KLTN.docx -> KLTN_clean.docx.

Khó: chú thích (Bold, size...), (Mẫu), dấu chấm ……… nằm CHUNG đoạn với nội dung
thật (vd "KHOA CÔNG NGHỆ THÔNG TIN (Bold, size 16)"). Nên với mỗi đoạn ta CẮT
phần chú thích trong đoạn (giữ nội dung + định dạng), chỉ XÓA hẳn đoạn khi sau
khi cắt nó rỗng hoặc là dòng ghi chú/ví dụ thuần túy.

Chạy: /tmp/thesis_venv/bin/python scripts/clean_thesis_template.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "thesis" / "KLTN.docx"
OUT = ROOT / "docs" / "thesis" / "KLTN_clean.docx"

# Chú thích trong ngoặc cần cắt bỏ (ngoặc có chứa các từ khóa này).
PAREN_ANNOT = re.compile(
    r"\([^)]*(?:bold|size|mẫu|in hoa|in đậm|được xếp|ký tên|đánh số|tùy theo|xếp sau)[^)]*\)",
    re.IGNORECASE,
)

# Đoạn là GHI CHÚ / VÍ DỤ thuần -> xóa hẳn.
NOTE_PREFIXES = [
    "ghi chú", "bold, size", "xếp sau trang", "số thứ nhất chỉ", "số thứ hai chỉ",
    "tên đề của bảng", "tên của hình", "ở cuối mỗi hình", "cụm từ viết",
    "danh mục tài liệu tham khảo tiếng", "danh mục tài liệu tham khảo trên",
    "danh mục tài liệu tham khảo xếp", "viết ngắn gọn", "nói rõ lý do",
    "in đậm và in hoa", "chữ số thứ", "nội dung size", "trình bày mỗi trang",
    "bắt đầu đánh số", "tên chương", "lời mở đầu:", "kết luận",
    "danh mục tài liệu tam khảo",
]
# Tên mẫu nằm CHUNG với nhãn (vd "Sinh viên thực hiện: TRẦN THỊ HOA") -> chỉ cắt tên.
INLINE_SAMPLES = ["trần thị hoa"]
# Dòng TLTK ví dụ -> xóa cả dòng.
SAMPLE_NAMES = ["nguyễn văn an", "nguyễn văn bằng", "an cư với lạc nghiệp", "chi tiết máy"]


def clean_text(text: str) -> str:
    text = PAREN_ANNOT.sub("", text)
    text = re.sub(r"\(\s*mẫu[^)]*\)", "", text, flags=re.IGNORECASE)
    for sample in INLINE_SAMPLES:
        text = re.sub(sample, "", text, flags=re.IGNORECASE)
    text = re.sub(r"[.…]{2,}", "", text)          # dòng/đuôi chấm
    text = re.sub(r"\t+\d*\s*$", "", text)             # số trang mẫu sau tab
    return re.sub(r"\s{2,}", " ", text).strip()


def is_pure_note(text: str) -> bool:
    low = text.lower()
    if low.startswith("-"):
        return True
    if any(low.startswith(prefix) for prefix in NOTE_PREFIXES):
        return True
    if any(name in low for name in SAMPLE_NAMES):
        return True
    # Mục lục/danh mục mẫu còn sót: "Chương 2", "1.1", "Bảng 1.1", "Hình 1.1"...
    if re.fullmatch(r"(chương\s*\d+|\d\.\d(\.\d)?|bảng\s*\d.*|hình\s*\d.*|….?|…)", low):
        return True
    return False


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        first = paragraph.runs[0]
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
        first.text = text
    else:
        paragraph.add_run(text)


def delete(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def main() -> None:
    document = Document(str(SRC))
    paragraphs = list(document.paragraphs)

    cover_indices = [i for i, p in enumerate(paragraphs) if "BỘ GIÁO DỤC VÀ ĐÀO TẠO" in p.text]
    first_cover = cover_indices[0]

    to_remove = []
    to_remove.extend(paragraphs[:first_cover])  # phần hướng dẫn trang 1-3

    if len(cover_indices) >= 2:  # trang bìa trùng lặp thứ 2
        start = cover_indices[1]
        end = next((i for i in range(start, len(paragraphs))
                    if "LỜI MỞ ĐẦU" in paragraphs[i].text), start)
        to_remove.extend(paragraphs[start:end])

    skip = {id(p) for p in to_remove}
    for paragraph in paragraphs:
        if id(paragraph) in skip:
            continue
        original = paragraph.text.strip()
        if not original:
            continue
        cleaned = clean_text(original)
        if not cleaned or is_pure_note(cleaned):
            to_remove.append(paragraph)
        elif cleaned != original:
            set_paragraph_text(paragraph, cleaned)

    for paragraph in to_remove:
        delete(paragraph)

    document.save(str(OUT))

    check = Document(str(OUT))
    print(f"Saved {OUT.name} | paragraphs còn lại: {len(check.paragraphs)}")
    print("=" * 60)
    for p in check.paragraphs:
        t = p.text.strip()
        if t:
            sizes = sorted({int(r.font.size.pt) for r in p.runs if r.font.size})
            print(f"[b={int(any(r.bold for r in p.runs))}|sz={sizes}] {t[:75]}")


if __name__ == "__main__":
    main()
