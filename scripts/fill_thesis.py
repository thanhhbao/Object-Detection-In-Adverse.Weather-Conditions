#!/usr/bin/env python3
"""Điền nội dung vào KLTN_clean.docx -> KLTN_full.docx.

Giữ định dạng file mẫu trường. Việc làm:
- Bìa: chèn tên đề tài (thông tin cá nhân để placeholder).
- Lời mở đầu / Lời cảm ơn: điền nội dung.
- Mục lục, Danh mục bảng/hình: chèn trường tự động (Word bấm F9 cập nhật).
- Bảng chữ viết tắt: điền.
- Chương 1-5: chèn trước trang PHỤ LỤC; tiêu đề Heading 1/2/3 để vào mục lục.
- Đánh số trang: La Mã (đầu) -> Ả Rập (từ Chương 1) bằng section break.

Chạy: /tmp/thesis_venv/bin/python scripts/fill_thesis.py
"""

from __future__ import annotations

import copy
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

# Tên style heading tự tạo (file mẫu .doc không có sẵn). outline level để mục
# lục tự động (TOC \o "1-3") nhận diện. level 1->0, 2->1, 3->2.
HEADING = {1: ("KL_Chuong", 16, 0), 2: ("KL_Muc", 14, 1), 3: ("KL_TieuMuc", 13, 2)}

sys.path.insert(0, str(Path(__file__).resolve().parent))
from thesis_content import ABBREVIATIONS, CHAPTERS, LOI_CAM_ON, LOI_MO_DAU

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "thesis" / "KLTN_clean.docx"
OUT = ROOT / "docs" / "thesis" / "KLTN_full.docx"

TITLE = ("THIẾT KẾ VÀ PHÁT TRIỂN MÔ HÌNH HỌC SÂU PHÁT HIỆN NGƯỜI VÀ PHƯƠNG TIỆN "
         "TRONG ĐIỀU KIỆN THỜI TIẾT BẤT LỢI PHỤC VỤ GIÁM SÁT GIAO THÔNG THÔNG MINH")

CENTER = WD_ALIGN_PARAGRAPH.CENTER
JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY


def find(document, predicate):
    for p in document.paragraphs:
        if predicate(p.text.strip()):
            return p
    return None


def find_last(document, predicate):
    found = None
    for p in document.paragraphs:
        if predicate(p.text.strip()):
            found = p
    return found


def new_para_after(ref_paragraph) -> Paragraph:
    element = OxmlElement("w:p")
    ref_paragraph._p.addnext(element)
    return Paragraph(element, ref_paragraph._parent)


def fmt(paragraph, *, size=13, bold=False, center=False, justify=False, indent=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
    if center:
        paragraph.alignment = CENTER
    elif justify:
        paragraph.alignment = JUSTIFY
    if indent:
        paragraph.paragraph_format.first_line_indent = Cm(1.0)
    paragraph.paragraph_format.line_spacing = 1.5
    return paragraph


def add_field(paragraph, instruction, placeholder):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = placeholder
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, sep, t, end):
        run._r.append(node)


def configure_headings(document):
    for level, (name, size, outline) in HEADING.items():
        try:
            style = document.styles[name]
        except KeyError:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        pPr = style.element.get_or_add_pPr()
        for existing in pPr.findall(qn("w:outlineLvl")):
            pPr.remove(existing)
        lvl = OxmlElement("w:outlineLvl"); lvl.set(qn("w:val"), str(outline))
        pPr.append(lvl)


def footer_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = CENTER
    run = para.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instr); run._r.append(end)


def set_numbering(sectPr, fmt_value, start=None):
    for existing in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(existing)
    pg = OxmlElement("w:pgNumType"); pg.set(qn("w:fmt"), fmt_value)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    sectPr.append(pg)


def build(document):
    configure_headings(document)

    # 1) Tên đề tài trên bìa (ngay sau "KHÓA LUẬN TỐT NGHIỆP").
    kltn = find(document, lambda t: t == "KHÓA LUẬN TỐT NGHIỆP")
    if kltn is not None:
        title = new_para_after(kltn)
        title.add_run(TITLE)
        fmt(title, size=18, bold=True, center=True)

    # 2) Lời mở đầu / Lời cảm ơn.
    for heading_text, paragraphs in (("LỜI MỞ ĐẦU", LOI_MO_DAU), ("LỜI CẢM ƠN", LOI_CAM_ON)):
        heading = find(document, lambda t, x=heading_text: t == x)
        after = heading
        for text in paragraphs:
            para = new_para_after(after)
            para.add_run(text)
            fmt(para, size=13, justify=True, indent=True)
            after = para

    # 3) Mục lục / Danh mục bảng / Danh mục hình -> trường tự động.
    for heading_text, instruction in (
        ("MỤC LỤC", 'TOC \\o "1-3" \\h \\z \\u'),
        ("DANH MỤC BẢNG, HÌNH", 'TOC \\h \\z \\c "Bảng"'),
        ("DANH MỤC HÌNH", 'TOC \\h \\z \\c "Hình"'),
    ):
        heading = find(document, lambda t, x=heading_text: t == x)
        if heading is not None:
            field_p = new_para_after(heading)
            add_field(field_p, instruction, "Bấm Ctrl+A rồi F9 để cập nhật.")
            fmt(field_p, size=13)

    # 4) Bảng chữ viết tắt.
    if document.tables:
        table = document.tables[0]
        for abbr, meaning in ABBREVIATIONS:
            cells = table.add_row().cells
            cells[0].text = abbr
            cells[1].text = meaning

    # 5) Chèn Chương 1-5 trước trang PHỤ LỤC (giữ thứ tự nhờ insert_paragraph_before).
    phu_luc = find_last(document, lambda t: t == "PHỤ LỤC")
    for chapter in CHAPTERS:
        h = phu_luc.insert_paragraph_before(chapter["title"], style=HEADING[1][0])
        h.alignment = CENTER
        for kind, text in chapter["body"]:
            if kind == "h2":
                phu_luc.insert_paragraph_before(text, style=HEADING[2][0])
            elif kind == "h3":
                phu_luc.insert_paragraph_before(text, style=HEADING[3][0])
            else:
                fmt(phu_luc.insert_paragraph_before(text), size=13, justify=True, indent=True)

    # 6) Section break: phần đầu La Mã -> từ Chương 1 Ả Rập.
    chuong1 = find(document, lambda t: t.startswith("CHƯƠNG 1"))
    prev = chuong1._p.getprevious()
    if prev is not None and prev.tag == qn("w:p"):
        pPr = prev.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr"); prev.insert(0, pPr)
        front_sectPr = copy.deepcopy(document.sections[-1]._sectPr)
        set_numbering(front_sectPr, "lowerRoman", start=1)
        pPr.append(front_sectPr)

    set_numbering(document.sections[-1]._sectPr, "decimal", start=1)
    for section in document.sections:
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        footer_page_number(section)


def main():
    document = Document(str(SRC))
    build(document)
    document.save(str(OUT))
    check = Document(str(OUT))
    print(f"Saved {OUT.name} | paragraphs: {len(check.paragraphs)} | sections: {len(check.sections)}")
    print("--- Chương (Heading 1) ---")
    for p in check.paragraphs:
        if p.style.name == HEADING[1][0]:
            print("  ", p.text[:60])


if __name__ == "__main__":
    main()
